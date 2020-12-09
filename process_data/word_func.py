# -*- coding: utf-8 -*-
"""
Created on Tue Nov 24 23:30:04 2020

@author: 61426
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import os
from tools import file_tools
import pandas as pd


# id date_list _10UV_list _10FG6_list
def generate_word_local(predict_df, save_path='./data/word'):
    # 存放路径
    print("预测数据存入本地"+Station_ID+".doc中")

    # nt = datetime.datetime.now()
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 文档标题
    title = document.add_paragraph()
    title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_cont = title.add_run('漳州市风速预测')
    title_cont.font.size = Pt(16)
    title_cont.bold = True
    pic = document.add_paragraph()
    # 图片居中设置
    pic.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pic_cont = pic.add_run("")
    pic_cont.add_picture(r'./test.png')#, width=Inches(2)

    # document.add_picture(r'./test.png')

    # 文档中插入表格 11行*4列        Table Grid
    table = document.add_table(rows=len(predict_df) + 1, cols=4, style='Medium Grid 1 Accent 1')
    table.autofit = False

    # 设置每列宽度
    table.columns[0].width = Cm(20)
    table.columns[1].width = Cm(20)
    table.columns[2].width = Cm(10)
    table.columns[3].width = Cm(10)

    ID_cols = table.columns[0].cells
    date_cols = table.columns[1].cells
    _10UV_cols = table.columns[2].cells
    _10FG6_cols = table.columns[3].cells
    ID_cols[0].add_paragraph('站点').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_cols[0].add_paragraph('日期').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _10UV_cols[0].add_paragraph('平均分').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _10FG6_cols[0].add_paragraph('阵风').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    ID_group = predict_df.groupby("id").groups
    ID_list = list(ID_group.keys())
    # ID date 10UV 10FG6
    merge_begin = 1
    merge_index = 1
    for ID in ID_list:
        ob_part = predict_df.loc[predict_df['id'] == ID].reset_index(drop=True)
        table.cell(len(ob_part) + merge_begin - 1, 0).merge(table.cell(merge_begin, 0))
        ID_cols[merge_begin].add_paragraph(ID).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        merge_begin = len(ob_part) + merge_begin
        for index in range(len(ob_part)):
            date_cols[merge_index].add_paragraph(ob_part.loc[index, 'date']).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _10UV_cols[merge_index].add_paragraph(ob_part.loc[index, '10UV']).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _10FG6_cols[merge_index].add_paragraph(ob_part.loc[index, '10FG6']).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            merge_index += 1
    file_tools.check_dir_and_mkdir(save_path)
    word_path = os.path.join(save_path, Station_ID + '.docx')
    document.save(word_path)


date_list = ['2015-08-01', '2015-08-02', '2015-08-03', '2015-08-04', '2015-08-05',
             '2015-08-06', '2015-08-07', '2015-08-08', '2015-08-09', '2015-08-10']

_10UV_list = ['6.1', '6.2', '6.3', '6.4', '6.5', '6.6', '6.7', '6.8', '6.9', '6.19']
_10FG6_list = ['5.1', '5.2', '5.3', '5.4', '5.5', '5.6', '5.7', '5.8', '5.9', '5.19']
Station_ID = 'F2273'

data_dict = {
    'id': ['F2273', 'F2273', 'F2273', 'F2274', 'F2274', 'F2274', 'F2275', 'F2275', 'F2275', 'F2276'],
    'date': date_list,
    '10UV': _10UV_list,
    '10FG6': _10FG6_list
}

# data_dict = {
#     'id': ['F2273', 'F2273', 'F2273', 'F2273', 'F2273', 'F2273', 'F2273', 'F2273', 'F2273', 'F2273'],
#     'date': date_list,
#     '10UV': _10UV_list,
#     '10FG6': _10FG6_list
# }
data_df = pd.DataFrame(data_dict)

generate_word_local(data_df)
